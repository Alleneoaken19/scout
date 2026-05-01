"""Resume builder -- PDF (WeasyPrint) + DOCX (python-docx) generation."""

import json
import shutil
import uuid
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader
from rich.console import Console
from weasyprint import HTML

from src.database import Resume, managed_session
from src.paths import RESUME_DIR

console = Console()

TEMPLATE_DIR = RESUME_DIR / "templates"
OUTPUT_DIR = RESUME_DIR / "generated"

_MONTH_NAMES = [
    "Jan", "Feb", "Mar", "Apr", "May", "Jun",
    "Jul", "Aug", "Sep", "Oct", "Nov", "Dec",
]


def format_date(value: str) -> str:
    """Convert '2022-08' to 'Aug 2022', 'present' to 'Present'."""
    if not value or not isinstance(value, str):
        return value or ""
    if value.lower() in ("present", "current"):
        return "Present"
    parts = value.split("-")
    if len(parts) == 2:
        try:
            year, month = parts
            return f"{_MONTH_NAMES[int(month) - 1]} {year}"
        except (ValueError, IndexError):
            return value
    return value


def _normalize_skills(skills):
    """Ensure skills is a dict of category -> list for template rendering."""
    if isinstance(skills, dict):
        return skills
    if isinstance(skills, list):
        return {"Skills": skills}
    return {}


@dataclass
class ResumeResult:
    id: str
    pdf_path: str
    docx_path: str


def build_pdf(tailored: dict, job_id: str) -> Path:
    """Render tailored resume data to PDF using Jinja2 + WeasyPrint."""
    out_dir = OUTPUT_DIR / job_id
    out_dir.mkdir(parents=True, exist_ok=True)

    env = Environment(  # nosec B701 - templates are internal, not user-uploaded
        loader=FileSystemLoader(str(TEMPLATE_DIR)),
        autoescape=True,
    )
    env.filters["format_date"] = format_date

    try:
        template = env.get_template("resume.html")
    except Exception as e:
        raise FileNotFoundError(f"Resume template not found at {TEMPLATE_DIR}/resume.html: {e}")

    render_data = {**tailored, "skills": _normalize_skills(tailored.get("skills", {}))}
    html_content = template.render(**render_data)

    pdf_path = out_dir / "resume.pdf"
    HTML(string=html_content).write_pdf(str(pdf_path))

    if not pdf_path.exists():
        raise RuntimeError(f"PDF generation failed -- file not created at {pdf_path}")

    return pdf_path


def build_docx(tailored: dict, job_id: str) -> Path:
    """Generate a DOCX resume from tailored data."""
    out_dir = OUTPUT_DIR / job_id
    out_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "Calibri"

    personal = tailored.get("personal", {})

    # Header
    heading = doc.add_heading(personal.get("name", ""), level=0)
    heading.style.font.size = Pt(18)

    contact_parts = [
        personal.get("location", ""),
        personal.get("email", ""),
        personal.get("phone", ""),
    ]
    doc.add_paragraph(" | ".join(p for p in contact_parts if p))

    # Summary
    summary = tailored.get("summary", "")
    if summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

    # Experience
    experience = tailored.get("experience", [])
    if experience:
        doc.add_heading("Experience", level=1)
        for exp in experience:
            role = exp.get("role", "Role")
            company = exp.get("company", "Company")
            start = format_date(exp.get("start_date", ""))
            end = format_date(exp.get("end_date", ""))
            date_range = f" ({start} \u2013 {end})" if start else ""
            doc.add_heading(f"{role} | {company}{date_range}", level=2)
            for bullet in exp.get("bullets", []):
                text = bullet if isinstance(bullet, str) else bullet.get("text", str(bullet))
                doc.add_paragraph(text, style="List Bullet")

    # Skills (categorized)
    skills = _normalize_skills(tailored.get("skills", {}))
    if skills:
        doc.add_heading("Skills", level=1)
        for category, items in skills.items():
            if items:
                p = doc.add_paragraph()
                run = p.add_run(f"{category}: ")
                run.bold = True
                p.add_run(", ".join(items))

    # Projects
    projects = tailored.get("projects", [])
    if projects:
        doc.add_heading("Projects", level=1)
        for proj in projects:
            name = proj.get("name", "") if isinstance(proj, dict) else str(proj)
            desc = proj.get("description", "") if isinstance(proj, dict) else ""
            doc.add_paragraph(f"{name}: {desc}" if desc else name, style="List Bullet")

    # Education
    education = tailored.get("education", [])
    if education:
        doc.add_heading("Education", level=1)
        for edu in education:
            degree = edu.get("degree", "") if isinstance(edu, dict) else str(edu)
            institution = edu.get("institution", "") if isinstance(edu, dict) else ""
            year = edu.get("year", "") if isinstance(edu, dict) else ""
            parts = [p for p in [degree, institution, f"({year})" if year else ""] if p]
            doc.add_paragraph(" \u2014 ".join(parts))

    docx_path = out_dir / "resume.docx"
    doc.save(str(docx_path))

    return docx_path


def build_resume(tailored: dict, job_id: str, ats_score_val: float) -> ResumeResult:
    """Build PDF + DOCX and save a Resume record to the database.

    If DB commit fails, cleans up generated files to avoid orphans.
    """
    out_dir = OUTPUT_DIR / job_id
    pdf_path = None
    docx_path = None

    try:
        pdf_path = build_pdf(tailored, job_id)
        docx_path = build_docx(tailored, job_id)

        resume_id = str(uuid.uuid4())
        with managed_session() as session:
            resume = Resume(
                id=resume_id,
                job_id=job_id,
                pdf_path=str(pdf_path),
                docx_path=str(docx_path),
                ats_score=ats_score_val,
                tailored_json=json.dumps(tailored),
                created_at=datetime.now(UTC),
            )
            session.add(resume)

        return ResumeResult(id=resume_id, pdf_path=str(pdf_path), docx_path=str(docx_path))

    except Exception:
        # Clean up orphaned files on failure
        if out_dir.exists():
            shutil.rmtree(out_dir, ignore_errors=True)
        raise
